# -*- coding: utf-8 -*-
"""生成 MES 概要设计、开发计划和需求规约三份 Word 文档。

外部依赖为 python-docx 和仓库根目录的 DOCX 模板。DOCX 本质是 ZIP + OOXML：脚本会处理
Strict/Transitional 命名空间兼容、目录域、表格样式、页眉页脚和隐私元数据，最终写入 docs/。
代码只生成文档，不连接 MES 数据库或调用业务接口。
"""

from __future__ import annotations

import os
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PROJECT = "双星轮胎制造MES系统"
VERSION = "V1.0"
DOC_DATE = "2026年07月16日"


def strict_to_transitional(source: Path) -> Path:
    """把 Strict OOXML 模板内的命名空间 URI 转为 python-docx 可读取的 Transitional URI。

    原 DOCX 不修改；转换后的临时 ZIP 写到系统临时目录，XML/关系文件逐项替换，其余媒体原样复制。
    """
    replacements = {
        b"http://purl.oclc.org/ooxml/officeDocument/relationships": b"http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        b"http://purl.oclc.org/ooxml/wordprocessingml/main": b"http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        b"http://purl.oclc.org/ooxml/officeDocument/math": b"http://schemas.openxmlformats.org/officeDocument/2006/math",
        b"http://purl.oclc.org/ooxml/drawingml/wordprocessingDrawing": b"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        b"http://purl.oclc.org/ooxml/drawingml/main": b"http://schemas.openxmlformats.org/drawingml/2006/main",
        b"http://purl.oclc.org/ooxml/drawingml/picture": b"http://schemas.openxmlformats.org/drawingml/2006/picture",
        b"http://purl.oclc.org/ooxml/drawingml/chart": b"http://schemas.openxmlformats.org/drawingml/2006/chart",
    }
    target = Path(tempfile.gettempdir()) / "mes-requirements-template-transitional.docx"
    with ZipFile(source) as zin, ZipFile(target, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith((".xml", ".rels")):
                for old, new in replacements.items():
                    data = data.replace(old, new)
            zout.writestr(item, data)
    return target


def scrub_personal_metadata(path: Path) -> None:
    """Clear author fields without opening Word, preserving the document body."""
    target = path.with_suffix(path.suffix + ".metadata-tmp")
    with ZipFile(path) as zin, ZipFile(target, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                from lxml import etree
                root = etree.fromstring(data)
                for element in root.xpath('//*[local-name()="creator" or local-name()="lastModifiedBy"]'):
                    element.text = ""
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    os.replace(target, path)


def clear_document_body(doc: Document) -> None:
    """删除模板正文但保留 sectPr 页面设置，以复用纸张、页边距、页眉和页脚。"""
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    """通过底层 OOXML 为表格单元格设置十六进制底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    """给表格首行添加 w:tblHeader，使跨页表格在每页重复表头。"""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9, center: bool = False) -> None:
    """重建单元格段落和 run，统一字号、粗体、对齐、字体及多行文本。"""
    cell.text = ""
    lines = str(text).split("\n") if text is not None else [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_doc(doc: Document, title: str, subject: str) -> None:
    """设置文档属性、A4 页面、页边距、默认字体、标题样式、页眉和自动页码。"""
    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "MES, 轮胎制造, 项目文档"

    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)

        section.header.is_linked_to_previous = False
        header = section.header
        header.paragraphs[0].text = PROJECT
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.paragraphs[0].runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 116, 139)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        section.footer.is_linked_to_previous = False
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{PROJECT}  ·  第 ")
        r.font.size = Pt(9)
        add_field(p, "PAGE")
        r = p.add_run(" 页")
        r.font.size = Pt(9)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)):
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(15, 45, 86)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    """插入 Word 域代码，例如 TOC 或 PAGE；打开 Word 更新域后生成目录/页码。"""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_cover(doc: Document, document_title: str, document_no: str) -> None:
    """创建统一封面，写入系统名、文档标题、编号、版本、日期和项目单位。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    run = p.add_run(PROJECT)
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(11, 61, 117)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run(document_title)
    run.bold = True
    run.font.size = Pt(32)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(VERSION)
    run.font.size = Pt(18)

    for _ in range(4):
        doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    labels = [("文档编号", document_no), ("编制", ""), ("审核", ""), ("批准", ""), ("日期", DOC_DATE)]
    for row, (label, value) in zip(table.rows, labels):
        set_cell_text(row.cells[0], label, bold=True, center=True, size=11)
        set_cell_text(row.cells[1], value, center=True, size=11)
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(8)
    doc.add_page_break()


def add_revision_history(doc: Document, content: str) -> None:
    """添加修改履历页和首次发布记录，随后插入分页符。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("修 改 履 历")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    table = add_table(
        doc,
        ["修改编号", "日期", "修改人", "版本号", "修改内容"],
        [["1", DOC_DATE, "", VERSION, content], ["", "", "", "", ""], ["", "", "", "", ""]],
        widths=[1.8, 2.8, 2.3, 2.0, 7.0],
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.add_page_break()


def add_toc(doc: Document, max_level: int = 3) -> None:
    """插入最多 max_level 级的自动目录域及分页符，目录内容由 Word 更新域生成。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p = doc.add_paragraph()
    add_field(p, f'TOC \\o "1-{max_level}" \\h \\z \\u', "打开文档后更新目录")
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1):
    """添加指定级别标题并设置与后段同页，返回段落供调用方继续扩展。"""
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc: Document, text: str, *, first_indent: bool = True, bold_prefix: str | None = None):
    """添加正文段落，可控制中文首行缩进并将指定前缀加粗。"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if first_indent else None
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str], *, numbered: bool = False) -> None:
    """添加统一悬挂缩进的项目符号或人工编号列表。"""
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.add_run(f"{index}. " if numbered else "• ").bold = numbered
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    """创建带网格、蓝色表头、重复标题行和可选列宽的标准项目文档表格。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, center=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc: Document, text: str) -> None:
    """添加居中的图表题注，并设置与相邻正文协调的段前段后距。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(71, 85, 105)


def build_overview_design() -> Path:
    """基于概要设计模板生成总体架构、模块、接口、数据、安全、部署与测试章节。"""
    template = ROOT / "项目概要设计书.docx"
    output = DOCS / "项目概要设计书.docx"
    doc = Document(template)
    clear_document_body(doc)
    format_doc(doc, f"{PROJECT}项目概要设计书", "项目概要设计")
    add_cover(doc, "项目概要设计书", "MES-OD-001")
    add_revision_history(doc, "依据当前前后端实现、数据库设计和权限方案完成概要设计。")
    add_toc(doc, 4)

    add_heading(doc, "文档概述", 1)
    add_heading(doc, "文档目的和范围", 2)
    add_body(doc, "本文档从总体结构、业务模块、关键业务时序、数据结构和接口边界五个方面描述双星轮胎制造MES系统，为详细设计、程序开发、联调测试、部署验收和后续维护提供统一依据。")
    add_body(doc, "设计范围覆盖Vue前端、JAX-RS后端、PostgreSQL数据库及其权限控制，包含计划与工单、生产报工、仓储物流、质量管理、设备维护、工艺与主数据、产品追溯、管理反馈、用户与权限，以及工作台和经营驾驶舱。")

    add_heading(doc, "术语/缩略语", 2)
    add_table(doc, ["序号", "术语/缩略语", "说明"], [
        ["1", "MES", "Manufacturing Execution System，制造执行系统。"],
        ["2", "PMC", "Production and Material Control，生产与物料控制。"],
        ["3", "BOM", "Bill of Materials，产品物料清单。"],
        ["4", "SOP", "Standard Operating Procedure，标准作业程序。"],
        ["5", "RBAC", "基于角色的访问控制，角色关联权限点，用户可拥有一个或多个角色。"],
        ["6", "数据范围", "在功能权限之外，限定用户可访问的产线、仓库或本人业务数据。"],
        ["7", "AGV/运输机器人", "执行拣货后物料配送的自动化设备，在系统中作为设备或接口对象，不作为人工登录角色。"],
        ["8", "JAX-RS", "Java RESTful Web Services规范，本系统后端HTTP接口的实现标准。"],
        ["9", "JDBC", "Java数据库连接接口，DAO层通过JDBC访问PostgreSQL。"],
    ], widths=[1.4, 3.4, 11.2])

    add_heading(doc, "参考文档", 2)
    add_table(doc, ["序号", "文档名", "来源", "日期/版本"], [
        ["1", "软件分析设计报告.docx", "项目组", "2026-05-26"],
        ["2", "双星轮胎制造MES系统数据库设计书.xlsx", "项目资料", "当前版"],
        ["3", "权限管理正式实施说明-v3.md", "项目资料", "v3"],
        ["4", "后端重构与接口审计.md", "项目资料", "当前版"],
        ["5", "module_table_mapping.md", "数据库设计输出", "当前版"],
        ["6", "前后端源代码及自动化测试", "项目仓库", "2026-07-16"],
    ], widths=[1.4, 6.8, 4.0, 3.8])

    add_heading(doc, "系统结构图", 1)
    add_body(doc, "系统采用前后端分离的浏览器/服务器结构。前端负责页面路由、会话状态、权限可见性和业务交互；后端按照Controller、Service、DAO、Entity分层；PostgreSQL保存业务数据、权限数据和审计数据。后端是最终鉴权边界。")
    add_table(doc, ["层次", "主要技术/组件", "职责"], [
        ["用户交互层", "Vue 3、Vue Router、Pinia、Lucide", "登录、导航、表格、表单、看板、移动端适配和公开追溯页面。"],
        ["接口层", "JAX-RS / Jersey、Jackson", "接收HTTP参数，调用业务服务，统一返回ApiResponse。"],
        ["业务层", "Java 21 Service", "校验业务规则、角色边界、数据范围、状态流转和跨模块用例编排。"],
        ["数据访问层", "DAO、JDBC、事务", "封装SQL、行映射、查询以及需原子提交的库存、工单等事务。"],
        ["数据层", "PostgreSQL", "保存订单、任务、工单、库存、报工、质量、设备、追溯、权限和审计数据。"],
        ["运行与部署", "Tomcat 10.1、WAR、Nginx（可选）", "托管REST接口和前端构建产物，支持反向代理与HTTPS。"],
    ], widths=[2.5, 4.2, 9.3])
    add_caption(doc, "图2-1 系统逻辑分层结构（以表格表示）")

    image_path = DOCS / "系统模块图.jpg"
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(6.35))
        add_caption(doc, "图2-2 双星轮胎制造MES系统总流程图")

    add_heading(doc, "部署结构", 2)
    add_bullets(doc, [
        "开发环境：Vite开发服务器运行Vue应用，并将/api请求代理到本地Java后端。",
        "生产环境：Vue构建产物打包进WAR，由Tomcat统一托管；Vue Router使用History回退到index.html。",
        "数据库连接：后端通过环境变量读取PostgreSQL连接信息，数据库迁移由专用迁移入口执行。",
        "公开入口：轮胎二维码指向公开追溯路由，公开接口仅返回允许展示的追溯数据。",
    ])

    add_heading(doc, "模块概述", 1)
    add_heading(doc, "模块功能定义", 2)
    module_rows = [
        ["1", "计划与工单", "客户订单、生产任务、齐套分析、缺料预警、制造工单、派工接单、返工重排", "PMC计划员、车间管理员、生产操作工"],
        ["2", "生产报工", "制造工单接收、报工提交与审核、合格/不良数量、工时、计件工资", "生产操作工、车间管理员、人事经理"],
        ["3", "仓储物流", "领料申请与审批、库存、拣货、运输机器人配送、采购入库、物料/仓库/库位", "仓库管理员、生产操作工"],
        ["4", "质量管理", "创建抽检、分配质检员、录入检验项、提交与判定、返工单和质量追溯", "质量主管、质检员"],
        ["5", "设备维护", "设备台账、故障报修、审核转维修、维修派工、结果提交、验收和维护计划", "设备管理员、设备维护员、现场人员"],
        ["6", "工艺与主数据", "产品、BOM关联、工艺路线、工序顺序、生产线及基础制造能力", "工艺工程师、PMC计划员"],
        ["7", "产品追溯", "订单—任务—工单—批次追溯、合格轮胎标签、二维码、打印记录、公开查询", "质量、仓储、管理层及外部查询者"],
        ["8", "管理反馈", "关联制造工单登记生产、质量、设备、物料、交期等异常并闭环", "总经理、相关业务负责人"],
        ["9", "用户与权限", "用户账号、角色、权限点、多角色、产线/仓库数据范围、账号与权限申请、会话维护", "系统管理员、人事经理"],
        ["10", "工作台与经营驾驶舱", "按角色汇总待办、产量、订单、质量、设备、库存和异常指标", "全部登录用户、总经理"],
    ]
    add_table(doc, ["序号", "模块", "功能点详细内容", "主要使用者"], module_rows, widths=[1.1, 3.0, 8.5, 3.4])

    add_heading(doc, "模块结构", 2)
    add_table(doc, ["模块组", "组成模块", "依赖关系"], [
        ["业务入口", "工作台、经营驾驶舱、个人资料", "读取各业务域聚合数据，不绕过业务服务修改源数据。"],
        ["计划执行", "计划与工单、生产报工", "订单生成任务；任务齐套后形成工单；工单报工反写实际产量和状态。"],
        ["物料保障", "仓储物流", "工单触发领料；审批后生成拣货与配送；交接时扣减库存并记录流水。"],
        ["质量与设备保障", "质量管理、设备维护", "报工触发质量检验；不合格进入返工；故障进入报修和维修闭环。"],
        ["制造基础", "工艺与主数据", "为订单、齐套、排产、工单和设备提供产品、工艺、产线基础数据。"],
        ["数据闭环", "产品追溯、管理反馈", "聚合订单、批次、物料、工序、质量和设备记录，支撑异常反馈。"],
        ["安全底座", "用户与权限", "登录会话、RBAC、数据范围和审计贯穿全部受保护接口。"],
    ], widths=[2.5, 4.2, 9.3])

    add_heading(doc, "模块动作时序", 2)
    sequences = [
        ("主生产闭环", "登录鉴权 → 创建客户订单 → 创建生产任务 → 执行齐套分析 → 制定制造工单 → 车间派工 → 操作工接单 → 发起领料 → 仓库接收/审批 → 拣货 → 配送到达/确认收料 → 生产报工 → 车间审核 → 创建质检 → 质检提交 → 质量判定 → 生成轮胎标签并入库 → 产品追溯/经营看板。"),
        ("缺料处理闭环", "齐套分析发现不足 → 形成缺口明细 → PMC发布缺料预警 → 仓库接收预警 → 外部采购或库存调整入库 → 再次齐套分析 → 任务变为READY → 制定制造工单。"),
        ("质量返工闭环", "质检员提交不合格/返工结论 → 质量主管判定 → 生成返工单和质量追溯 → PMC将返工需求纳入计划 → 生成新任务/工单 → 重新生产、报工和质检。"),
        ("设备维修闭环", "现场人员报修 → 设备管理员审核 → 转维修工单 → 派给设备维护员 → 维护员提交结果 → 设备管理员验收 → 设备恢复可用状态 → 受影响生产继续执行。"),
        ("权限变更闭环", "人事经理发起账号/权限申请 → 系统管理员审核 → 执行角色或账号变更 → 更新会话和数据范围 → 记录申请、审核和关键操作。"),
    ]
    for name, content in sequences:
        add_heading(doc, name, 3)
        add_body(doc, content, first_indent=False)

    add_heading(doc, "接口说明", 1)
    add_heading(doc, "数据结构定义", 2)
    add_body(doc, "系统接口统一使用JSON，标准响应结构为ApiResponse<T>：success表示操作是否成功，message返回中文结果信息，data承载对象、列表或聚合结果。日期时间按ISO-8601字符串交换，数量与金额使用数值类型。")
    add_table(doc, ["业务对象", "核心表", "关键字段/关系"], [
        ["用户与权限", "mes_user、mes_role、mes_permission、mes_user_role、mes_role_permission、mes_role_data_scope、mes_user_session", "user_id、role_code、permission_code；用户—角色—权限及产线/仓库数据范围。"],
        ["客户订单", "mes_customer_order", "order_id、order_no、product_id、order_qty、delivery_date、order_status。"],
        ["生产任务与齐套", "mes_production_task、mes_kitting_analysis、mes_kitting_shortage_item、mes_shortage_alert", "task_id关联order_id；齐套结果和物料/产线/设备/工序缺口。"],
        ["制造工单", "mes_work_order、mes_work_order_operation_log", "work_order_id关联task_id、line_id、process_id和assigned_to；记录派工与状态日志。"],
        ["领料与库存", "mes_material_requisition、mes_material_requisition_item、mes_inventory、mes_inventory_transaction", "requisition_id关联work_order_id；库存按物料、仓库、库位、批次和质量状态管理。"],
        ["拣货与配送", "mes_picking_task、mes_robot、mes_robot_delivery_task", "拣货任务由已审批领料单生成；配送确认收料时扣减库存。"],
        ["报工与计件", "mes_work_report、mes_piecework_wage", "report_id关联work_order_id和operator_id；审核通过后更新工单实绩并生成工资记录。"],
        ["质量与返工", "mes_quality_inspection、mes_quality_inspection_item、mes_rework_order、mes_quality_trace", "质检关联报工和工单；判定结果驱动返工与质量追溯。"],
        ["设备维护", "mes_equipment、mes_equipment_repair_report、mes_maintenance_order、mes_maintenance_plan", "equipment_id贯穿台账、报修、维修派工、结果与验收。"],
        ["主数据与追溯", "mes_product、mes_product_bom、mes_process_route、mes_production_line、mes_product_trace、轮胎标签相关表", "产品、物料、工艺、产线支撑排产；追溯链关联订单、任务、工单、批次、质量和入库。"],
    ], widths=[2.4, 6.2, 7.4])

    add_heading(doc, "函数", 2)
    add_heading(doc, "模块间接口函数", 3)
    add_table(doc, ["接口组", "主要端点", "方法", "功能/权限边界"], [
        ["认证", "/api/auth/login、/me、/logout；/api/profile", "POST/GET/PUT", "登录签发会话；读取当前用户；退出和个人资料维护。"],
        ["计划工单", "/api/orders、/production-tasks、/work-orders、/shortage-alerts、/planning/reworks", "GET/POST", "订单、任务、齐套、预警、工单派接和返工重排；受planning.*权限控制。"],
        ["生产报工", "/api/work-reports、/piecework-wages", "GET/POST/PUT/DELETE", "报工提交、本人修改、审核驳回、计件查询；操作工与车间审核职责分离。"],
        ["仓储物流", "/api/requisitions、/inventory、/picking-tasks、/robot-delivery-tasks", "GET/POST/PUT/DELETE", "领料审批、库存调整、拣货配送和收料扣库；仓库和工单数据范围生效。"],
        ["质量", "/api/quality-inspections、/rework-orders、/quality-traces", "GET/POST", "创建/分配/检验/判定、返工执行和追溯；质检员与质量主管分权。"],
        ["设备", "/api/equipment、/equipment-repair-reports、/maintenance-orders、/maintenance-plans", "GET/POST/PUT", "设备、报修、维修派工、执行、验收和保养计划。"],
        ["主数据", "/api/products、/process-routes、/production-lines、/materials、/warehouses", "GET/POST/PUT/DELETE", "产品、工艺、产线、物料和仓库主数据维护。"],
        ["追溯", "/api/product-traces、/tire-labels、/public/tire-traces/{token}", "GET/POST", "内部追溯、合格轮胎标签及公开扫码；公开端点不要求人员会话。"],
        ["看板反馈", "/api/dashboard/my-summary、/executive、/management-feedback", "GET/POST", "按角色聚合指标和管理反馈闭环。"],
        ["权限管理", "/api/users、/access/roles、/permissions、/permission-applications、/account-applications", "GET/POST/PUT", "用户、角色、数据范围、账号与权限申请审批。"],
    ], widths=[2.2, 6.2, 2.5, 5.1])

    add_heading(doc, "模块内接口函数", 3)
    add_table(doc, ["服务/方法", "输入", "输出", "核心处理与事务边界"], [
        ["AuthService.login", "用户名、密码、客户端信息", "会话令牌、用户、角色、权限", "校验账号状态与密码哈希，创建可撤销会话并更新最近登录信息。"],
        ["PlanningService.analyzeKitting", "生产任务ID", "齐套分析与缺口明细", "综合BOM、库存、产线、设备和工序条件，更新任务齐套状态。"],
        ["WorkOrderService.create/dispatch/receive", "任务、产线、工序、数量、操作工", "制造工单/状态结果", "校验任务READY、产线和工序；原子写入工单与操作日志。"],
        ["WarehouseService.approveRequisition", "领料申请ID", "审批结果、拣货任务", "锁定并检查库存，审批成功后创建拣货任务；任一明细不足则整体失败。"],
        ["WarehouseService.confirmReceipt", "配送任务ID", "库存流水和完成状态", "原子扣减批次库存、写流水、完成领料明细和配送任务。"],
        ["ProductionService.approveReport", "报工ID", "工单实绩、计件记录", "审核报工，生成计件工资，累加工单合格数量并判断RUNNING/FINISHED。"],
        ["QualityInspectionService.judge", "质检ID、判定", "质检状态、返工/追溯结果", "审核检验项；不合格或返工时生成返工单并维护追溯链。"],
        ["EquipmentService.acceptMaintenance", "维修工单ID", "验收结果、设备状态", "校验维修已完成，验收后更新维修单和设备状态。"],
        ["TireLabelService.generate", "合格质检、仓库、库位、数量", "轮胎序列号、二维码、入库记录", "仅允许已审核合格质检，批量生成唯一标签并关联生产追溯。"],
        ["DataScopeService", "当前用户、资源类型", "SQL过滤条件/允许的ID集合", "在业务查询和操作前应用本人、产线和仓库范围，系统管理员按策略处理。"],
    ], widths=[4.0, 3.5, 3.5, 5.0])

    doc.save(output)
    return output


def set_plan_cell(table, row: int, col: int, text: str, *, bold: bool = False, center: bool = False, size: float = 8.5) -> None:
    """开发计划表的单元格便捷包装，统一调用 set_cell_text。"""
    set_cell_text(table.cell(row, col), text, bold=bold, center=center, size=size)


def build_development_plan() -> Path:
    """生成项目范围、组织分工、里程碑、WBS、风险、质量和配置管理计划。"""
    template = ROOT / "项目开发计划书.docx"
    output = DOCS / "项目开发计划书.docx"
    doc = Document(template)
    format_doc(doc, f"{PROJECT}项目开发计划书", "项目开发计划")
    if doc.paragraphs:
        doc.paragraphs[0].text = "项目开发计划书"
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in doc.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    table = doc.tables[0]
    set_plan_cell(table, 1, 2, PROJECT, bold=True, center=True, size=10)
    set_plan_cell(table, 1, 9, "", center=True)
    for row, col in ((2, 2), (2, 9), (3, 2), (3, 9), (4, 2), (4, 9)):
        set_plan_cell(table, row, col, "", center=True)
    for row in (5, 6, 7):
        set_plan_cell(table, row, 2, "")

    set_plan_cell(table, 8, 2, "2026年07月08日", center=True)
    set_plan_cell(table, 8, 8, "2026年07月31日", center=True)
    descriptions = [
        "建设面向轮胎制造现场的Web版MES，打通客户订单、生产任务、齐套分析、制造工单、领料配送、生产报工、质量检验、设备维修、产品追溯和经营看板。",
        "计划与工单：维护客户订单，创建生产任务，执行物料/产线/设备/工序齐套分析；缺料时发布预警，齐套后制定工单，并完成派工、接单与返工重排。",
        "生产与仓储：操作工按本人制造工单发起领料和报工；仓库完成领料接收、库存校验、审批、拣货、机器人配送与收料扣库；报工审核后形成计件记录并更新工单进度。",
        "质量与设备：质量主管创建和分配质检任务，质检员录入检验项并提交，主管判定后形成返工或合格入库；设备故障按报修、审核、维修派工、执行和验收闭环。",
        "追溯与管理：为合格轮胎生成唯一序列号和二维码，形成订单—任务—工单—批次—质量—入库追溯；工作台、经营驾驶舱和管理反馈支撑日常执行与决策。",
    ]
    # Rows 10-14 are one vertically merged cell in the supplied template.
    set_plan_cell(table, 10, 0, "\n".join(descriptions), size=8.2)

    set_plan_cell(table, 16, 2, "迭代式开发（需求—设计—编码—联调—验证）", center=True)
    set_plan_cell(table, 16, 9, "Java、JavaScript、SQL", center=True)
    set_plan_cell(table, 17, 2, "Java 21、JAX-RS/Jersey、JDBC、PostgreSQL、Tomcat 10.1、Vue 3、Vite、Pinia、Vue Router、JUnit 5、Node.js测试脚本", size=8.2)

    milestones = [
        ["第一里程碑", "07月08日", "07月11日", "4", "需求梳理、角色边界、总体架构和计划确认", "需求规约、概要设计、开发计划"],
        ["第二里程碑", "07月12日", "07月17日", "6", "数据库、RBAC和核心后端业务闭环开发", "数据库脚本、实体/DAO/Service/接口、后端测试"],
        ["第三里程碑", "07月18日", "07月24日", "7", "Vue前端、角色工作台、全流程联调与数据修正", "前端页面、接口联调记录、演示数据"],
        ["第四里程碑", "07月25日", "07月31日", "7", "系统测试、部署、文档完善和项目验收", "WAR包、测试记录、部署说明、验收材料"],
    ]
    for row_index, data in zip(range(20, 24), milestones):
        for col, value in zip((0, 1, 3, 5, 7, 9), data):
            set_plan_cell(table, row_index, col, value, center=col != 7, size=8)

    progress = [
        (26, "07月08日", "完成范围、角色和交付物确认；形成启动纪要。"),
        (27, "07月08日—07月10日", "完成业务流程、功能需求、非功能需求和验收标准评审。"),
        (28, "07月10日—07月11日", "完成登录、导航、九大模块、看板和移动端原型评审。"),
        (29, "07月11日", "统一Vue—JAX-RS—Service—DAO—PostgreSQL分层和接口规范。"),
        (30, "07月12日—07月17日", "完成数据库、接口、状态机、RBAC与数据范围详细设计。"),
        (31, "07月17日", "评审表结构、事务边界、权限矩阵和跨模块依赖。"),
        (32, "07月12日—07月24日", "按模块迭代实现，主分支集成前执行编译和自动化测试。"),
        (33, "每次合并前", "检查分层约束、接口唯一性、状态流转、异常处理和命名规范。"),
        (34, "07月20日—07月28日", "完成单元、接口、角色权限、全流程、构建和移动端冒烟测试。"),
        (35, "07月29日—07月30日", "整理成果物、问题清单、风险复盘和项目总结。"),
        (36, "07月31日", "部署演示环境，按主生产闭环和异常闭环进行答辩。"),
    ]
    for row, planned, result in progress:
        set_plan_cell(table, row, 1, planned, center=True, size=7.8)
        set_plan_cell(table, row, 3, "", center=True)
        set_plan_cell(table, row, 5, "", center=True)
        set_plan_cell(table, row, 7, result, size=7.8)

    document_schedule = [
        (39, "V1.0.0", "07月11日"),
        (40, "V1.0.0", "07月17日"),
        (41, "V1.0.0", "07月17日"),
        (42, "V1.0.0", "07月31日"),
    ]
    for row, version, planned in document_schedule:
        set_plan_cell(table, row, 1, version, center=True)
        set_plan_cell(table, row, 2, planned, center=True)
        set_plan_cell(table, row, 4, "", center=True)
        set_plan_cell(table, row, 7, "", center=True)
        set_plan_cell(table, row, 9, "待提交/待评审", center=True, size=8)

    measures = {
        45: "每日更新任务状态和阻塞项；提交前执行最小相关测试；跨模块接口当天确认请求、响应和状态约束。",
        46: "每周按里程碑核对计划与实际；集中评审数据库迁移、权限矩阵、接口变更和遗留缺陷；必要时调整优先级。",
        47: "每个里程碑以可运行成果物、测试记录和评审结论作为完成条件；未达标项进入下一阶段风险清单并指定关闭日期。",
    }
    for row, value in measures.items():
        set_plan_cell(table, row, 4, value, size=8)

    set_plan_cell(table, 50, 0, "需求或权限边界频繁变化，可能导致前后端接口和角色菜单反复调整。", size=8)
    set_plan_cell(table, 50, 4, "冻结版本范围；以权限矩阵和接口清单为基线；变更先评估影响，再同步代码、测试和文档。", size=8)
    set_plan_cell(table, 51, 0, "库存扣减、报工审核、质检返工等跨模块状态不一致；远程数据库测试可能污染演示数据。", size=8)
    set_plan_cell(table, 51, 4, "关键动作采用数据库事务与幂等校验；保留操作日志；默认运行不写远程库测试，集成测试使用独立数据并清理。", size=8)
    set_plan_cell(table, 53, 1, "具备Java、Vue和PostgreSQL基础；通过分层规范、代码评审、接口契约和自动化测试降低能力差异。", size=8)
    set_plan_cell(table, 54, 1, "复杂", center=True, bold=True)
    set_plan_cell(table, 55, 1, "", center=True)
    set_plan_cell(table, 55, 7, "", center=True)

    doc.save(output)
    return output


def add_requirement_detail(doc: Document, number: str, req_id: str, name: str, priority: str,
                           background: str, functions: str, constraints: str, data: str,
                           outputs: str, exceptions: str, acceptance: str) -> None:
    """以固定字段表添加一条详细需求，覆盖背景、功能、约束、数据、输出、异常和验收标准。"""
    add_heading(doc, f"{number} {name}", 3)
    rows = [
        ["需求编号", req_id],
        ["功能名称", name],
        ["优先级", priority],
        ["业务背景", background],
        ["功能说明", functions],
        ["约束条件", constraints],
        ["输入/数据", data],
        ["输出/状态", outputs],
        ["异常处理", exceptions],
        ["验收标准", acceptance],
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, center=True, size=9)
        set_cell_shading(cells[0], "EAF2F8")
        set_cell_text(cells[1], value, size=9)
        cells[0].width = Cm(3.0)
        cells[1].width = Cm(13.0)
    doc.add_paragraph()


def build_requirements() -> Path:
    """转换并读取需求规约模板，生成角色、功能、非功能、接口、数据与验收需求。"""
    source = ROOT / "需求规约.docx"
    template = strict_to_transitional(source)
    output = DOCS / "需求规约.docx"
    doc = Document(template)
    clear_document_body(doc)
    format_doc(doc, f"{PROJECT}需求规约", "软件需求规格说明")
    add_cover(doc, "需求规约", "MES-SRS-001")
    add_revision_history(doc, "依据当前系统范围、代码实现、数据库和权限方案形成需求基线。")
    add_toc(doc, 3)

    add_heading(doc, "引言", 1)
    add_heading(doc, "1.1 目的", 2)
    add_body(doc, "本文定义双星轮胎制造MES系统的业务目标、功能需求、角色边界、数据约束、安全要求、运行环境和验收标准，是设计、开发、测试、部署和项目验收的共同基线。")
    add_body(doc, "本文面向项目负责人、开发人员、测试人员、数据库人员、实施运维人员及业务验收人员。未在本文确认的个人姓名、学号、班级、组号、审批人等信息均不填写。")

    add_heading(doc, "1.2 背景", 2)
    add_body(doc, "轮胎制造涉及订单、排产、物料齐套、车间执行、仓储配送、质量检验、设备保障和成品追溯等多个环节。依赖人工表单或孤立系统会造成状态不一致、反馈滞后和追溯困难。本项目通过统一的制造执行平台串联现场业务，使生产任务、工单、库存、质量和设备信息可查询、可控制、可追溯。")
    add_body(doc, "系统名称为“双星轮胎制造MES系统”。系统采用Web方式访问，前端为Vue应用，后端为Java REST服务，数据库为PostgreSQL。外部ERP/WMS、机器人和设备目前作为边界或集成对象，人员账号与设备/系统凭证分离。")

    add_heading(doc, "1.3 参考资料", 2)
    add_bullets(doc, [
        "《软件分析设计报告.docx》，双星轮胎制造MES系统分析与设计资料。",
        "《双星轮胎制造MES系统数据库设计书.xlsx》及数据库建表、迁移脚本。",
        "《权限管理正式实施说明-v3.md》及角色权限验收资料。",
        "《后端重构与接口审计.md》、前后端源代码和自动化测试。",
        "《MES全功能端到端演练操作手册.md》及系统模块总流程图。",
    ])

    add_heading(doc, "1.4 术语", 2)
    add_table(doc, ["术语", "说明"], [
        ["客户订单", "客户对轮胎产品、数量和交付日期的需求，是生产任务来源。"],
        ["生产任务", "PMC基于客户订单形成的计划任务，需经过齐套分析后才能制定工单。"],
        ["制造工单", "下达至指定产线、工序和操作工的现场执行单元。"],
        ["齐套分析", "对BOM物料、库存、产线、设备和工艺条件进行综合检查。"],
        ["领料申请", "生产操作工针对制造工单申请生产物料的业务单据。"],
        ["生产报工", "操作工提交的产量、合格量、不良量和工时记录。"],
        ["质检/返工", "对报工结果进行检验、审核和判定；不合格时形成返工闭环。"],
        ["追溯码/轮胎二维码", "关联订单、任务、工单、批次、质量和入库信息的唯一标识。"],
        ["RBAC", "基于角色和权限点的访问控制模型，并叠加产线、仓库和本人数据范围。"],
    ], widths=[4.0, 12.0])

    add_heading(doc, "任务概述", 1)
    add_heading(doc, "2.1 目标", 2)
    add_bullets(doc, [
        "建立订单到轮胎成品的统一数据链，减少重复录入和信息断点。",
        "实现生产任务齐套检查、制造工单派接、领料配送、报工审核、质检返工和设备维修闭环。",
        "实现按角色授权、按产线/仓库/本人限定数据范围，确保审批与执行职责分离。",
        "为每条合格轮胎生成可打印二维码，支持内部和公开的全流程追溯。",
        "通过角色工作台和经营驾驶舱提供待办、产量、质量、设备、库存及异常决策信息。",
        "提供可部署、可测试、可迁移和可维护的前后端分层工程。",
    ], numbered=True)

    add_heading(doc, "2.2 用户与角色", 2)
    add_table(doc, ["角色编码", "角色名称", "主要职责"], [
        ["SYSTEM_ADMIN", "系统管理员", "用户、角色、权限、数据范围、会话和系统维护；关键动作仍需留痕。"],
        ["HR_MANAGER", "人事经理", "人员信息、账号申请、权限变更申请和人员维度报工查询。"],
        ["GENERAL_MANAGER", "总经理/管理层", "经营驾驶舱、追溯、异常反馈和决策查询。"],
        ["PMC_PLANNER", "PMC计划员", "订单、生产任务、齐套分析、缺料预警、制造工单和返工重排。"],
        ["WORKSHOP_MANAGER", "车间管理员", "工单派发、生产报工审核、现场进度与异常协调。"],
        ["PRODUCTION_OPERATOR", "生产操作工", "接收本人工单、发起领料、确认收料、提交和维护本人报工。"],
        ["WAREHOUSE_ADMIN", "仓库管理员", "领料接收审批、库存、拣货、配送、采购入库及仓储主数据。"],
        ["QUALITY_MANAGER", "质量主管", "创建质检、分配质检员、质量判定和返工管理。"],
        ["QUALITY_INSPECTOR", "质检员", "执行已分配检验、录入检验项并提交结果。"],
        ["PROCESS_ENGINEER", "工艺工程师", "产品、BOM、工艺路线、工序与制造参数维护。"],
        ["EQUIPMENT_ADMIN", "设备管理员", "设备台账、报修审核、维修派工、验收和维护计划。"],
        ["EQUIPMENT_MAINTAINER", "设备维护员", "接收维修工单、执行维修并提交结果。"],
    ], widths=[3.4, 3.4, 9.2])

    add_heading(doc, "2.3 系统边界", 2)
    add_bullets(doc, [
        "系统内：人员登录、业务单据、状态流转、库存事务、质量/设备闭环、追溯、看板、权限和审计。",
        "系统外：客户、ERP/WMS正式生产接口、财务工资结算、设备PLC/传感器、机器人厂商控制系统和短信/企业微信等通知通道。",
        "机器人与外部系统不使用普通员工账号；后续对接时应采用独立API客户端凭证和最小权限。",
    ])

    add_heading(doc, "需求规定", 1)
    add_heading(doc, "3.1 一般性需求", 2)
    add_table(doc, ["编号", "类别", "需求"], [
        ["GR-01", "一致性", "相同业务对象在列表、详情、看板和追溯中使用统一编码、中文标签和状态含义。"],
        ["GR-02", "易用性", "业务表单优先通过下拉关联真实数据，减少要求用户直接输入数据库ID。"],
        ["GR-03", "正确性", "数量、状态、归属和前置条件在后端校验；前端校验仅用于及时提示。"],
        ["GR-04", "可靠性", "库存扣减、领料审批、报工审核、质检返工等关键动作应在事务中原子完成。"],
        ["GR-05", "性能", "教学/演示规模下普通列表和业务操作应在3秒内返回；3万条轮胎追溯数据查询应在5秒内给出结果。"],
        ["GR-06", "兼容性", "支持主流Chromium浏览器，桌面端1366×768及以上可完整操作，390px宽度可完成核心流程。"],
        ["GR-07", "可维护性", "后端遵循Controller→Service→DAO→PostgreSQL分层；数据库变更以版本化SQL迁移交付。"],
        ["GR-08", "可观测性", "接口返回统一成功标识与中文错误信息；登录、权限和高风险业务动作可审计。"],
        ["GR-09", "数据保护", "生产配置不得在源码中硬编码真实数据库口令；演示账号与生产账号隔离。"],
    ], widths=[2.0, 2.5, 11.5])

    add_heading(doc, "3.2 功能性需求", 2)
    add_body(doc, "系统功能模块如下表所示。模块能否显示由用户角色和权限决定；任何前端隐藏均不能替代后端授权校验。")
    add_table(doc, ["功能模块", "主要功能", "裁剪说明"], [
        ["登录、账号与权限", "登录、退出、个人资料、用户、角色、权限、数据范围、账号/权限申请和会话维护", "不可裁剪"],
        ["计划与工单", "订单、生产任务、齐套、缺料预警、制造工单、派接和返工重排", "不可裁剪"],
        ["生产报工", "工单接收、报工、审核、计件工资", "不可裁剪"],
        ["仓储物流", "领料、审批、库存、拣货、配送、采购入库和仓储主数据", "不可裁剪"],
        ["质量管理", "质检任务、检验项、提交判定、返工和质量追溯", "不可裁剪"],
        ["设备维护", "设备、报修、维修派工、执行验收和维护计划", "不可裁剪"],
        ["工艺与主数据", "产品、BOM、工艺路线、生产线", "不可裁剪"],
        ["产品追溯", "追溯链、轮胎二维码、标签打印、公开查询", "不可裁剪"],
        ["工作台、经营驾驶舱与管理反馈", "角色待办、经营指标和工单异常闭环", "驾驶舱可按部署范围裁剪"],
    ], widths=[4.0, 9.2, 2.8])

    image_path = DOCS / "系统模块图.jpg"
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(6.35))
        add_caption(doc, "图3-1 系统功能与整体业务流程")

    details = [
        ("3.2.1", "FR-AUTH", "登录、账号与权限管理", "高",
         "系统有12类正式人员角色，各岗位必须使用独立账号登录，并按照职责和数据范围访问功能。",
         "1）用户名/密码登录、读取当前会话、退出；2）维护个人资料；3）系统管理员维护用户、角色、权限点、多角色和产线/仓库数据范围；4）人事经理发起账号或权限申请，系统管理员审核/执行；5）维护删除账号记录、锁定会话和会话撤销。",
         "用户名唯一；密码只保存哈希；禁用、删除、锁定或会话失效的账号不得访问；前端菜单与后端权限点保持一致；未匹配的受保护接口默认拒绝。",
         "username、password、realName、roleCodes、department、phone、lineIds、warehouseIds、申请原因与审核意见。",
         "会话令牌、当前用户、角色/权限集合；账号和申请状态；审计与会话记录。",
         "凭据错误、账号禁用/锁定、无权限、越权数据访问、目标角色无效时返回明确错误且不改变数据。",
         "12类验收账号均可登录并看到正确模块；无权限接口返回拒绝；修改角色或撤销会话后访问结果立即符合新权限。"),
        ("3.2.2", "FR-PLAN", "计划与制造工单", "高",
         "PMC需要将客户订单转换为可执行生产任务，并在物料、产线、设备和工艺齐套后下达制造工单。",
         "1）创建/查询客户订单；2）创建生产任务；3）执行齐套分析并显示缺口；4）发布和接收缺料预警；5）选择产线、工序、数量和批次制定制造工单；6）车间派工、操作工接单；7）记录工单操作日志；8）将质量返工纳入新生产计划。",
         "订单和产品必须存在；任务数量合法；仅READY且齐套的任务可制定工单；产线可用、工序匹配产品；仅CREATED工单可派发，且只能由被派操作工接单。",
         "客户、产品、数量、交期、优先级、任务、产线、工序、计划数量、批次、操作工。",
         "订单/任务/工单编号；齐套结果和缺口；工单状态CREATED→DISPATCHED→RECEIVED→RUNNING→FINISHED。",
         "缺料、产线故障、工序不匹配、数量超限、重复派工/接单或越权时拒绝操作并保留原状态。",
         "能够从待排产订单创建任务；缺料时不能制定工单且可发布预警；补料后齐套成功并可完成派工接单。"),
        ("3.2.3", "FR-PROD", "生产报工与计件", "高",
         "生产操作工需要按本人已接收工单报告产量和工时，车间管理员审核后更新工单实绩并形成计件记录。",
         "1）查看/接收本人制造工单；2）提交报工；3）在允许状态下修改本人报工；4）车间管理员审核通过或驳回；5）查询本人或授权范围计件工资。",
         "工单状态为RECEIVED或RUNNING；操作工必须是工单被派人员；合格量+不良量应与报工量一致；累计报工不得超过计划量允许范围；同一报工只审核一次。",
         "workOrderId、reportQty、qualifiedQty、defectQty、workHours、驳回原因。",
         "报工单状态；审核通过后生成计件记录、累加工单actualQty并更新RUNNING/FINISHED。",
         "数量不一致、工单未接收、不是本人、超报、重复审核或状态非法时不写入或不更新。",
         "操作工只能看到并操作本人数据；审核通过后工单实绩和计件金额同步变化；驳回原因可见。"),
        ("3.2.4", "FR-WH", "仓储物流", "高",
         "生产需要按制造工单申领物料，仓库完成库存核对、审批、拣货、机器人配送和现场交接。",
         "1）操作工发起领料并查看本人/工单申请；2）仓库接收、批准或驳回；3）审批生成拣货任务；4）完成拣货生成配送任务；5）记录到达并由现场确认收料；6）确认时扣减库存并写流水；7）维护物料、仓库、库位、机器人和库存；8）执行外部采购入库。",
         "领料必须关联有效且归属当前操作工的执行中工单；允许缺料时创建申请，但批准时逐项检查合格可用库存；收料扣库必须原子执行且不可重复。",
         "工单、仓库、物料明细、需求数量、单位、批次、库位、采购数量与原因。",
         "领料状态CREATED→RECEIVED→APPROVED→COMPLETED或REJECTED；拣货/配送任务；库存余额与流水。",
         "库存不足时审批失败且不生成拣货；批次/仓库/库位无效、重复交接或越权时拒绝；事务失败全部回滚。",
         "完整执行领料—审批—拣货—配送—收料后库存准确减少且有流水；缺料申请能创建但不能错误批准。"),
        ("3.2.5", "FR-QA", "质量检验与返工", "高",
         "经审核的生产报工必须进行质量检验，质量职责需分离，并对不合格产品建立返工和追溯。",
         "1）质量主管从待检报工创建质检单并分配质检员；2）质检员录入项目标准值、实测值和项目结论；3）提交总体结果；4）质量主管审核判定；5）不合格/返工时生成返工单；6）派发、完成返工并维护质量追溯。",
         "质检必须关联有效报工；质检员只能执行分配给自己的任务；至少有检验项才能提交；只有SUBMITTED状态可判定；返工不得重复生成。",
         "workReportId、inspectionType、inspectorId、检验项目、标准值、实际值、PASS/FAIL/REWORK及备注。",
         "质检状态CREATED→IN_PROGRESS→SUBMITTED→REVIEWED；判定结果；返工单和质量追溯记录。",
         "无待检报工、未分配、无检验项、状态非法、重复判定或越权时拒绝并提示。",
         "质量主管和质检员权限分离；合格可进入轮胎标签流程；不合格能形成返工并由PMC重新排产。"),
        ("3.2.6", "FR-EQP", "设备维护", "高",
         "生产设备故障会影响排产与执行，需要从报修到验收的闭环，并保留设备台账和计划维护信息。",
         "1）设备台账和状态维护；2）现场发起故障报修；3）设备管理员审核并转维修工单；4）派给维护员；5）维护员提交维修结果；6）设备管理员验收；7）维护周期计划维护。",
         "设备编码唯一；报修设备必须有效；未审核报修不得转维修；维护员只能处理分配任务；未完成维修不得验收。",
         "设备、关联工单、故障级别/描述、维护员、维修结果、维护周期和下次维护时间。",
         "报修状态、维修工单CREATED→ASSIGNED/IN_PROGRESS→FINISHED→ACCEPTED、设备RUNNING/IDLE/FAULT/MAINTENANCE。",
         "设备无效、重复转单、维护员不匹配、状态非法或越权时拒绝；验收失败不恢复可用状态。",
         "故障设备可完成报修、审核、派工、维修和验收；验收后设备状态正确并可重新用于齐套/排产。"),
        ("3.2.7", "FR-MASTER", "工艺与主数据", "高",
         "订单、齐套、排产和工单依赖稳定的产品、BOM、工艺路线、生产线和仓储主数据。",
         "1）维护产品编码、名称、型号和规格；2）维护产品BOM及用料；3）按产品维护工艺路线、工序顺序和工作中心；4）维护生产线类型、产能和状态；5）仓储侧维护物料、仓库和库位。",
         "各类业务编码唯一；停用主数据不可用于新业务；工序顺序不得冲突；产线故障/停用时不可排产；已被业务引用的数据不得物理删除造成断链。",
         "产品、BOM物料及用量、工序编码/名称/顺序/工作中心、产线及产能、物料/仓库/库位。",
         "可供订单、齐套分析、工单、领料和设备模块选择的有效主数据。",
         "编码重复、引用不存在、顺序冲突、状态不允许或删除受引用数据时拒绝。",
         "新增主数据后能在相关表单中选择；停用后不再出现在新业务候选项；工艺路线按正确顺序展示。"),
        ("3.2.8", "FR-TRACE", "产品追溯与轮胎二维码", "高",
         "轮胎产品需要以唯一序列号关联生产和质量履历，支持内部查询和消费者/客户扫码。",
         "1）创建订单—任务—工单—批次追溯主链；2）从已审核合格质检单批量生成轮胎序列号和标签；3）记录打印次数；4）提供二维码、标签和追溯文档；5）公开令牌查询允许展示的追溯信息。",
         "仅REVIEWED且PASS的质检可生成标签；序列号、追溯码和公开令牌唯一；生成数量不得超过可入库合格数量；公开接口不得暴露内部账号、权限和敏感信息。",
         "inspectionId、workOrderId、warehouseId、locationId、quantity、打印备注、公开token。",
         "轮胎序列号、二维码图像、标签/追溯文档、入库关联和公开追溯视图。",
         "质检不合格、数量超限、重复生成、仓库无效或token不存在时返回明确结果且不生成脏数据。",
         "合格质检可生成唯一标签并打印；扫码能查询到产品、批次、关键工序和质量信息；无效token不泄露数据。"),
        ("3.2.9", "FR-DASH", "工作台、经营驾驶舱与管理反馈", "中",
         "不同岗位需要快速查看待办和业务指标，管理层需要了解订单、产量、质量、设备、库存和异常并形成反馈。",
         "1）按当前角色返回工作台汇总与待办；2）总经理查看经营驾驶舱；3）按权限查看产量、订单、质量、设备和库存指标；4）对制造工单创建管理反馈并关闭；5）反馈按工单数据范围查询。",
         "指标必须来自业务源数据或受控聚合，不允许客户端自行拼接改变口径；非总经理不得访问专属驾驶舱；反馈关闭需相应权限。",
         "当前用户、时间范围、工单、反馈类型和内容。",
         "角色摘要、经营指标、异常列表、反馈OPEN/CLOSED状态。",
         "无数据时返回零值/空列表；无权限、无工单归属或重复关闭时拒绝。",
         "12类角色登录后均有合理工作台；总经理可查看驾驶舱；反馈能创建、查询、关闭并受数据范围保护。"),
    ]
    for detail in details:
        add_requirement_detail(doc, *detail)

    add_heading(doc, "3.3 系统安全性的要求", 2)
    add_heading(doc, "3.3.1 数据存储安全", 3)
    add_bullets(doc, [
        "密码只保存不可逆哈希，不在接口、日志和数据库中保存或返回明文密码。",
        "数据库账号、口令和外部服务密钥通过环境变量配置；示例配置不包含真实生产秘密。",
        "库存、工单、报工、质量和维修关键动作使用事务，失败时回滚，避免部分成功。",
        "部署环境应定期备份PostgreSQL，并验证恢复流程；迁移前对关键数据进行备份。",
    ])
    add_heading(doc, "3.3.2 访问控制安全", 3)
    add_bullets(doc, [
        "除登录和公开追溯外，接口必须校验有效会话。",
        "后端根据HTTP方法和规范化路径匹配权限点，受保护接口未明确授权时默认拒绝。",
        "在功能权限上叠加本人、产线和仓库数据范围，防止同角色跨范围访问。",
        "审批与执行分离：报工提交/审核、质检执行/判定、维修执行/验收、权限申请/执行由不同职责完成。",
        "禁用账号、删除账号、撤销会话和锁定会话应立即阻止后续访问。",
    ])
    add_heading(doc, "3.3.3 网络传输安全", 3)
    add_bullets(doc, [
        "生产部署使用HTTPS，HTTP应重定向到HTTPS；反向代理限制请求体大小和异常方法。",
        "会话令牌仅通过Authorization头发送，不写入公开URL；公开追溯只使用随机、不可推测的token。",
        "跨域策略仅允许受信任来源；错误响应不得回显数据库口令、SQL或服务器绝对路径。",
    ])
    add_heading(doc, "3.3.4 应用系统审计", 3)
    add_bullets(doc, [
        "记录登录、退出、登录失败、账号/角色/数据范围变更、会话撤销等安全事件。",
        "记录库存调整与扣减、报工审核、质量判定、返工、维修验收、删除和管理反馈关闭等高风险操作。",
        "审计记录至少包含操作者、时间、客户端、动作、资源、资源ID、结果和必要变更摘要。",
        "普通业务角色不得修改或删除审计记录。",
    ])
    add_heading(doc, "3.3.5 系统约束", 3)
    add_bullets(doc, [
        "后端运行于Java 21，使用Jakarta/Jersey API；部署容器兼容Servlet 6。",
        "数据库采用PostgreSQL，SQL迁移需保持向前兼容并可在演示环境重复验证。",
        "Controller不得直接访问DAO/JDBC，Service不得直接执行SQL，DAO不得依赖JAX-RS。",
        "机器人和外部系统不得复用人员登录账号；正式集成应使用独立API客户端与最小权限。",
    ])
    add_heading(doc, "3.3.6 其他专门要求", 3)
    add_bullets(doc, [
        "中文界面中的状态、角色、字段名和错误信息应使用业务可理解的中文标签。",
        "公开追溯页面兼顾手机扫码访问，并对不存在、过期或禁用token提供友好提示。",
        "任何演示初始密码仅限验收环境，首次正式使用前必须修改。",
    ])

    add_heading(doc, "运行环境规定", 1)
    add_heading(doc, "4.1 运行环境", 2)
    add_heading(doc, "4.1.1 软件环境", 3)
    add_table(doc, ["类别", "要求"], [
        ["服务器操作系统", "Windows 10/11或主流64位Linux发行版。"],
        ["Java", "JDK 21。"],
        ["应用容器", "Tomcat 10.1.x；生产可由Nginx提供反向代理与HTTPS。"],
        ["数据库", "PostgreSQL 14或以上（建议），PostgreSQL JDBC 42.7.4。"],
        ["前端构建", "Node.js 20或以上、Vue 3.5、Vite 7、Pinia 3、Vue Router 4。"],
        ["客户端", "主流Chromium内核浏览器；启用JavaScript、Cookie/本地会话存储。"],
        ["测试", "JUnit 5、Node.js测试脚本；浏览器冒烟测试使用Playwright Core。"],
    ], widths=[4.0, 12.0])
    add_heading(doc, "4.1.2 硬件环境", 3)
    add_table(doc, ["场景", "最低建议"], [
        ["开发机", "4核CPU、8GB内存、10GB可用磁盘、可访问PostgreSQL。"],
        ["演示/实训服务器", "2核CPU、4GB内存、20GB可用磁盘；数据库可同机或独立部署。"],
        ["客户端", "1366×768及以上桌面显示器；核心流程兼容390px宽移动端。"],
        ["网络", "客户端到服务器网络稳定；生产部署建议内网访问并提供HTTPS出口。"],
    ], widths=[4.0, 12.0])

    add_heading(doc, "遗留问题", 1)
    add_table(doc, ["编号", "遗留项", "影响与后续处理"], [
        ["OI-01", "ERP/WMS正式接口", "当前以本系统订单、库存和同步日志为主；需与真实外部系统确认字段、频率、重试和对账机制。"],
        ["OI-02", "设备/机器人实时接入", "当前完成设备、机器人和配送业务建模；需根据实际PLC、AGV厂商协议实现采集与控制适配器。"],
        ["OI-03", "通知渠道", "当前业务待办可在系统内查看；短信、邮件、企业微信等外部通知需另行选型和配置。"],
        ["OI-04", "高级排程与智能建议", "当前支持规则化齐套和计划建议；若上线智能排程，需补充算法目标、训练/规则数据和可解释性验收。"],
        ["OI-05", "生产级容量指标", "本文性能指标适用于实训/演示规模；正式上线前需根据用户数、数据量和硬件完成压力测试。"],
    ], widths=[2.0, 4.2, 9.8])

    add_heading(doc, "项目非技术需求", 1)
    add_table(doc, ["编号", "类别", "要求"], [
        ["NTR-01", "交付物", "交付源代码、数据库脚本、构建产物、需求规约、设计文档、开发计划、部署说明和测试记录。"],
        ["NTR-02", "文档", "业务状态、角色权限、接口和数据库字段变更时同步更新相应文档。"],
        ["NTR-03", "培训", "为系统管理员和各业务角色提供演练账号、操作手册和主流程演示。"],
        ["NTR-04", "验收", "按主生产闭环、缺料闭环、质量返工闭环、设备维修闭环和权限边界进行验收。"],
        ["NTR-05", "质量", "合并前通过编译、自动化测试、接口唯一性检查、分层约束测试和前端生产构建。"],
        ["NTR-06", "维护", "数据库变更以增量迁移执行，环境配置独立于代码，禁止通过删除历史业务数据解决兼容问题。"],
        ["NTR-07", "隐私", "个人姓名、电话等只向具有业务必要性的角色展示；文档中无法确认的个人信息保持空白。"],
    ], widths=[2.0, 2.7, 11.3])

    doc.save(output)
    try:
        os.unlink(template)
    except OSError:
        pass
    return output


def main() -> None:
    """创建输出目录，顺序生成三份文档，清除作者元数据并打印最终路径。"""
    DOCS.mkdir(exist_ok=True)
    outputs = [build_overview_design(), build_development_plan(), build_requirements()]
    for path in outputs:
        scrub_personal_metadata(path)
        print(path)


if __name__ == "__main__":
    main()
